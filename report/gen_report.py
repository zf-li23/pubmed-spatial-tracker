#!/usr/bin/env python3
"""Generate monthly report charts and .docx for zf-li23."""
import sqlite3, os, sys
from collections import Counter
from datetime import datetime

# ── matplotlib setup ──
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import numpy as np

# Try to find a Chinese font
import matplotlib.font_manager as fm
for f in fm.fontManager.ttflist:
    if 'Hei' in f.name or 'Song' in f.name or 'Ming' in f.name:
        plt.rcParams['font.sans-serif'] = [f.name, 'DejaVu Sans']
        break
else:
    plt.rcParams['font.sans-serif'] = ['DejaVu Sans']
plt.rcParams['axes.unicode_minus'] = False

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUT = os.path.dirname(os.path.abspath(__file__))
DB = os.path.join(os.path.dirname(OUT), 'spatial_literature.db')
CHARTS = OUT  # same dir

# ═══════════════════════════════════════
# DATABASE QUERIES
# ═══════════════════════════════════════
conn = sqlite3.connect(DB)

# Category distribution
cur = conn.execute("SELECT category, COUNT(*) FROM literature GROUP BY category ORDER BY COUNT(*) DESC")
cat_data = [(r[0] if r[0] else 'Uncategorized', r[1]) for r in cur.fetchall()]
total_articles = sum(s for _, s in cat_data)

# Year distribution
cur = conn.execute("SELECT pub_year, COUNT(*) FROM literature WHERE pub_year != '' AND pub_year IS NOT NULL GROUP BY pub_year ORDER BY pub_year")
year_data = [(int(r[0]), r[1]) for r in cur.fetchall() if r[0].isdigit() and 2010 <= int(r[0]) <= 2026]

# Tags top 25
cur = conn.execute("SELECT tags FROM literature WHERE tags IS NOT NULL AND tags != ''")
tc = Counter()
for (t,) in cur.fetchall():
    for tag in str(t).split(';'):
        tag = tag.strip()
        if tag and tag != 'nan' and tag != 'Discarded':
            tc[tag] += 1
top_tags = tc.most_common(25)

# Discarded count
cur = conn.execute("SELECT COUNT(*) FROM literature WHERE tags LIKE '%Discarded%'")
discarded_cnt = cur.fetchone()[0]

# Confirmed vs unconfirmed
cur = conn.execute("SELECT is_manually_confirmed, COUNT(*) FROM literature GROUP BY is_manually_confirmed")
conf_data = {r[0]: r[1] for r in cur.fetchall()}
confirmed = conf_data.get(1, 0)
unconfirmed = conf_data.get(0, 0)

# Journals top 10
cur = conn.execute("SELECT journal, COUNT(*) FROM literature WHERE journal IS NOT NULL AND journal != '' GROUP BY journal ORDER BY COUNT(*) DESC LIMIT 10")
journal_data = [(r[0][:60], r[1]) for r in cur.fetchall()]

# Uncertainty by category
cur = conn.execute("SELECT category, AVG(CAST(uncertainty_score AS REAL)), COUNT(*) FROM literature WHERE uncertainty_score IS NOT NULL AND uncertainty_score != '' GROUP BY category")
unc_data = {r[0]: (round(r[1], 2), r[2]) for r in cur.fetchall()}

# Uncertainty overall
cur = conn.execute("SELECT CAST(uncertainty_score AS REAL) FROM literature WHERE uncertainty_score IS NOT NULL AND uncertainty_score != ''")
u_scores = [r[0] for r in cur.fetchall()]

# MeSH top 15
cur = conn.execute("SELECT mesh_terms FROM literature WHERE mesh_terms IS NOT NULL AND mesh_terms != ''")
mesh_c = Counter()
for (m,) in cur.fetchall():
    for term in str(m).split(';'):
        term = term.strip()
        if term and term != 'nan':
            mesh_c[term] += 1
top_mesh = mesh_c.most_common(15)

# Category-year for trends
cur = conn.execute("SELECT category, pub_year, COUNT(*) FROM literature WHERE pub_year != '' AND category != '' GROUP BY category, pub_year ORDER BY pub_year")
cat_year = {}
for cat, yr, cnt in cur.fetchall():
    if yr.isdigit() and 2020 <= int(yr) <= 2026:
        cat_year.setdefault(cat, {})[int(yr)] = cnt

# Naive vs manual accuracy sample
cur = conn.execute("""SELECT category, naive_category, COUNT(*) 
                      FROM literature 
                      WHERE is_manually_confirmed=1 AND category!='' AND naive_category!='' 
                      GROUP BY category, naive_category""")
naive_vs = cur.fetchall()

# Preprint stats
cur = conn.execute("SELECT is_preprint, COUNT(*) FROM literature GROUP BY is_preprint")
preprint_data = {r[0]: r[1] for r in cur.fetchall()}

conn.close()

print(f"Total articles: {total_articles}")
print(f"Confirmed: {confirmed} / {total_articles} ({100*confirmed/total_articles:.2f}%)")
print(f"Top tag: {top_tags[0] if top_tags else 'N/A'}")

# ═══════════════════════════════════════
# CHARTS
# ═══════════════════════════════════════

# Chart 1: Yearly distribution (bar chart)
fig, ax = plt.subplots(figsize=(8, 4.5))
x, y = zip(*year_data)
colors = ['#3498db']*len(x)
# Highlight 2025-2026
for i in range(len(x)):
    if x[i] >= 2025:
        colors[i] = '#e74c3c'
bars = ax.bar(range(len(x)), y, color=colors, edgecolor='white', linewidth=0.5)
ax.set_xticks(range(len(x)))
ax.set_xticklabels(x, rotation=45, fontsize=7)
ax.set_ylabel('Number of Publications', fontsize=10)
ax.set_title('Annual Publication Count in Spatial Transcriptomics (2010–2026)', fontsize=12, fontweight='bold')
for i, v in enumerate(y):
    if v >= 500:
        ax.text(i, v + max(y)*0.015, str(v), ha='center', fontsize=8, fontweight='bold')
ax.set_ylim(0, max(y)*1.12)
ax.spines['top'].set_visible(False)
ax.spines['right'].set_visible(False)
fig.tight_layout()
fig.savefig(os.path.join(CHARTS, 'chart_yearly.png'), dpi=150)
plt.close()
print("Chart 1/6: yearly done")

# Chart 2: Category pie
fig, ax = plt.subplots(figsize=(6, 5))
labels = [c for c, s in cat_data if s >= 30]
sizes = [s for c, s in cat_data if s >= 30]
pie_colors = ['#2ecc71','#3498db','#9b59b6','#e67e22','#e74c3c','#95a5a6']
wedges, texts, autotexts = ax.pie(sizes, labels=None, autopct='%1.1f%%', colors=pie_colors,
                                   startangle=140, pctdistance=0.8)
for at in autotexts: at.set_fontsize(9)
ax.legend(wedges, [f'{l} ({s})' for l, s in zip(labels, sizes)], title='Category',
          loc='center left', bbox_to_anchor=(1, 0.5), fontsize=8)
ax.set_title('Literature Category Distribution', fontsize=13, fontweight='bold')
fig.tight_layout()
fig.savefig(os.path.join(CHARTS, 'chart_category_pie.png'), dpi=150)
plt.close()
print("Chart 2/6: category pie done")

# Chart 3: Top 15 tags
fig, ax = plt.subplots(figsize=(7, 5))
tags_15 = top_tags[:15]
tx, ty = zip(*tags_15)
tx = tx[::-1]; ty = ty[::-1]
bars = ax.barh(range(len(tx)), ty, color='#e74c3c', height=0.65, edgecolor='white')
for i, (t, v) in enumerate(zip(tx, ty)):
    ax.text(v + max(ty)*0.012, i, f' {v}', va='center', fontsize=8)
ax.set_yticks(range(len(tx)))
ax.set_yticklabels(tx, fontsize=9)
ax.set_xlabel('Count', fontsize=10)
ax.set_title('Top 15 Tags in Library', fontsize=12, fontweight='bold')
ax.invert_yaxis()
ax.spines['top'].set_visible(False)
ax.spines['right'].set_visible(False)
fig.tight_layout()
fig.savefig(os.path.join(CHARTS, 'chart_tags.png'), dpi=150)
plt.close()
print("Chart 3/6: tags done")

# Chart 4: Uncertainty by category
fig, ax = plt.subplots(figsize=(6, 4.5))
cats_order = ['Review','Technology','Database','Research','Data Analysis','Discard']
cats_u = [c for c in cats_order if c in unc_data and c != 'Discard']
vals_u = [unc_data[c][0] for c in cats_u if c in unc_data]
bar_colors = ['#e74c3c' if v > 0.45 else '#f39c12' if v > 0.35 else '#2ecc71' for v in vals_u]
bars = ax.bar(range(len(cats_u)), vals_u, color=bar_colors, edgecolor='white', linewidth=0.8)
for i, (bar, val) in enumerate(zip(bars, vals_u)):
    ax.text(i, val + 0.012, str(val), ha='center', fontsize=10, fontweight='bold')
ax.set_xticks(range(len(cats_u)))
ax.set_xticklabels(cats_u, fontsize=9)
ax.set_ylabel('Mean Uncertainty Score', fontsize=10)
ax.set_title('Model Uncertainty by Category', fontsize=12, fontweight='bold')
ax.set_ylim(0, max(vals_u)*1.2)
ax.spines['top'].set_visible(False)
ax.spines['right'].set_visible(False)
fig.tight_layout()
fig.savefig(os.path.join(CHARTS, 'chart_uncertainty.png'), dpi=150)
plt.close()
print("Chart 4/6: uncertainty done")

# Chart 5: Confirmation status pie
fig, ax = plt.subplots(figsize=(5, 4.5))
ax.pie([confirmed, unconfirmed], labels=[f'Confirmed\n({confirmed})', f'Unconfirmed\n({unconfirmed})'],
       autopct='%1.2f%%', colors=['#2ecc71','#bdc3c7'], startangle=90, explode=(0.05, 0),
       textprops={'fontsize': 9})
ax.set_title('Manual Confirmation Status', fontsize=12, fontweight='bold')
fig.tight_layout()
fig.savefig(os.path.join(CHARTS, 'chart_confirmation.png'), dpi=150)
plt.close()
print("Chart 5/6: confirmation done")

# Chart 6: Category trends 2020-2026
fig, ax = plt.subplots(figsize=(8, 4.5))
years_range = list(range(2020, 2027))
cat_list = ['Research','Data Analysis','Review','Technology','Database']
cmap = ['#2ecc71','#3498db','#9b59b6','#e67e22','#e74c3c']
bottom = np.zeros(len(years_range))
for ci, cat in enumerate(cat_list):
    vals = [cat_year.get(cat, {}).get(y, 0) for y in years_range]
    ax.bar(years_range, vals, bottom=bottom, label=cat, color=cmap[ci], alpha=0.85, edgecolor='white', linewidth=0.3)
    bottom += np.array(vals)
ax.legend(loc='upper left', fontsize=8, ncol=3)
ax.set_ylabel('Publications', fontsize=10)
ax.set_title('Category-wise Publication Trends (2020–2026)', fontsize=12, fontweight='bold')
ax.spines['top'].set_visible(False)
ax.spines['right'].set_visible(False)
fig.tight_layout()
fig.savefig(os.path.join(CHARTS, 'chart_trends.png'), dpi=150)
plt.close()
print("Chart 6/6: trends done")

print("\nAll charts saved to", CHARTS)

# ═══════════════════════════════════════
# DOCX REPORT
# ═══════════════════════════════════════
doc = Document()

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.25

# Helper functions
def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    return h

def add_para(text, bold=False, italic=False, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    # Data
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
    # Col widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table

def add_image_placeholder(path, caption, width_inches=5.5):
    """Add image if exists, otherwise placeholder."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(path):
        run = p.add_run()
        run.add_picture(path, width=Inches(width_inches))
    else:
        run = p.add_run(f'[ 图表占位: {os.path.basename(path)} — 请手动插入 ]')
        run.font.size = Pt(9)
        run.italic = True
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.font.size = Pt(8.5)
    r.italic = True
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# ═══════════════════ CONTENT ═══════════════════

# ── TITLE PAGE ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('科研训练月度报告')
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('(第6份)')
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()

# Info table
info = [
    ['姓名 / 学号', '李哲夫 / 2023011400', '班级', '生32'],
    ['课题名称', '亚细胞水平空间转录组共定位关系检测与数据预处理方法开发', '', ''],
    ['工作日期', '2026.03.21 – 2026.04.17', '周均工时', '20 小时'],
    ['所在实验室', '杨雪瑞实验室', '导师', ''],
]
t = doc.add_table(rows=len(info), cols=4)
t.style = 'Light Shading Accent 1'
for i, row in enumerate(info):
    for j, val in enumerate(row):
        t.rows[i].cells[j].text = val
        for p in t.rows[i].cells[j].paragraphs:
            for run in p.runs:
                run.font.size = Pt(9.5)
                if j == 0:
                    run.bold = True

doc.add_page_break()

# ═══════════════════ SECTIONS ═══════════════════

# ── Section 1: 本月工作计划 ──
add_heading('一、本月工作计划', level=1)

add_para('本月核心目标为搭建并稳定运行 PubMed Spatial Tracker 文献管理平台，同时以文献研读为主线建立对空间转录组学领域的系统性认知。')

add_para('双线推进策略：', bold=True)
doc.add_paragraph('平台建设线：完成文献检索入库→规则初分类→人工标注→PDF归档→模型重训→不确定性排序的全链路闭环。', style='List Bullet')
doc.add_paragraph('文献研读线：投入主要时间阅读空间转录组学各子方向的关键文献，梳理技术路线、分析方法与生物问题的对应关系。', style='List Bullet')

doc.add_paragraph()

# ── Section 2: 本月完成情况 ──
add_heading('二、本月实验完成情况', level=1)

# 2.1 数据与平台建设
add_heading('2.1 数据与平台建设进展', level=2)

add_para(f'截至 2026 年 4 月 17 日，平台已积累并结构化存储文献 {total_articles:,} 篇，包含 22 个字段（PMID、标题、摘要、DOI、期刊、年份、类别、标签、人工确认状态、PDF 路径、MeSH 词表、引用次数等）。')

add_para('核心量化指标：', bold=True)

summary_data = [
    ['文献总量', f'{total_articles:,}'],
    ['人工确认样本', f'{confirmed:,}'],
    ['人工确认率', f'{confirmed/total_articles*100:.1f}%'],
    ['有效年份覆盖', f'2010–2026 ({total_articles - sum(1 for _,c in cat_data if c=="Uncategorized"):,} 条有年份记录)'],
    ['预印本数量', f'{preprint_data.get("1",0):,}'],
    ['Discarded 样本', f'{discarded_cnt:,}'],
    ['平均不确定性分数', f'{sum(u_scores)/len(u_scores):.3f}'],
]
add_table(['指标', '数值'], summary_data, col_widths=[7, 7])

doc.add_paragraph()

# Category table
add_para('类别分布：', bold=True)
cat_rows = [[c, f'{s:,}', f'{s/total_articles*100:.1f}%'] for c, s in cat_data]
add_table(['类别 (Category)', '数量', '占比'], cat_rows, col_widths=[5, 3, 3])

doc.add_paragraph()

# High-frequency tags
add_para('高频标签 (Top 20)：', bold=True)
tag_rows = [[t, f'{c:,}'] for t, c in top_tags[:20]]
add_table(['标签', '频次'], tag_rows, col_widths=[6, 4])

doc.add_paragraph()

# Journals
add_para('主要来源期刊 (Top 10)：', bold=True)
jn_rows = [[j, f'{c:,}'] for j, c in journal_data]
add_table(['期刊', '篇数'], jn_rows, col_widths=[9, 3])

doc.add_page_break()

# Charts section
add_heading('2.2 可视化分析', level=2)

add_image_placeholder(
    os.path.join(CHARTS, 'chart_yearly.png'),
    '图1: 2010–2026年空间转录组学文献年度发文量（2025-2026年为高亮红色）。该领域呈指数增长态势，2025年全年文献量已达2,638篇，2026年截至4月已达1,246篇。',
    5.2
)

add_image_placeholder(
    os.path.join(CHARTS, 'chart_category_pie.png'),
    '图2: 文献类别构成。Research类占比最高（54.5%），其次为Data Analysis（23.3%）和Review（18.0%）。Technology与Database类别合计仅占4.3%，是后续精细化标注的重点。',
    4.8
)

add_image_placeholder(
    os.path.join(CHARTS, 'chart_trends.png'),
    '图3: 2020–2026年各类别增长趋势。所有类别均呈上升趋势，Research类增长最快（2025年1,390篇），Data Analysis类次之（2025年628篇）。',
    5.2
)

add_image_placeholder(
    os.path.join(CHARTS, 'chart_tags.png'),
    '图4: 标签频次Top 15。Visium（2,878）和Cancer（2,821）显著领先，反映空间转录组学目前以10x Visium为主要技术平台、以肿瘤为主要应用场景。Pathology（1,745）和Clustering（1,185）分别反映组织病理学和聚类分析是核心分析任务。',
    5.0
)

add_image_placeholder(
    os.path.join(CHARTS, 'chart_confirmation.png'),
    '图5: 人工确认进度。当前确认率仅6.84%（481/7,029），大量样本仍依赖自动预测，标注效率是下月推进重点。',
    4.0
)

add_image_placeholder(
    os.path.join(CHARTS, 'chart_uncertainty.png'),
    '图6: 各类别模型不确定性分布。Review类不确定性最高（0.52），反映出综述类文献的语义边界模糊；Data Analysis类最低（0.34），方法学论文的文本特征较易学习。',
    4.8
)

doc.add_page_break()

# 2.3 System features
add_heading('2.3 系统功能与工程完成度', level=2)

add_para('本月已完成以下可复现功能模块：')

features = [
    ('文献抓取与增量更新', '基于 PubMed Entrez API，按动态检索策略（基础空间转录组通用词 + tags.json 所有技术标签）增量拉取，自动去重，手动导入 PMID 列可复现追溯。'),
    ('规则初始化与 ML 双轨推断', 'Naive 规则引擎（migrate_naive.py）通过关键词匹配冷启动；ML 管线（ml_pipeline.py）基于 TF-IDF + SVC 在人工确认样本上训练，输出类别预测与不确定性分数。两条链路的分类器约束策略（enforce_category_tag_policy）已对齐。'),
    ('人工标注闭环', '支持类别修正（5大类）、标签修改（多选+自定义）、Discarded 负样本标记、确认状态入库，前端提交后即时更新列表排序。'),
    ('PDF 工作流', '支持本地上传（按 Category/Tag/DOI 自动归档）、URL 抓取下载、仅保存外链三种模式，与标注信息同步写入数据库。'),
    ('模型重训与不确定性排序', '前端一键触发 ML pipeline 重训，基于已确认样本重新拟合后预测全部未确认文献，按不确定性分数降序排列，引导优先标注最不确定的样本。'),
    ('标签管理体系', 'tags.json 集中管理标签本体（domain/technology/analysis/metaCategory），前端 TagManager 支持增删改，后端同步更新数据库中的已有标注。'),
    ('GitHub Pages 静态部署', '支持 VITE_STATIC_ONLY 模式，纯前端静态只读部署，数据快照自动生成。'),
]
for title, desc in features:
    p = doc.add_paragraph()
    r1 = p.add_run(f'{title}：')
    r1.bold = True
    r1.font.size = Pt(10)
    r2 = p.add_run(desc)
    r2.font.size = Pt(10)

doc.add_paragraph()

# 2.4 文献阅读与认知推进
add_heading('2.4 文献阅读与认知推进', level=2)

add_para('本月科研时间的核心投入在于空间转录组学文献的系统性阅读与知识框架梳理。通过 PubMed Spatial Tracker 筛选和结构化，重点阅读了以下方向的文献：')

add_para('（1）技术平台对比', bold=True)
add_para('通过阅读 Visium、MERFISH、Stereo-seq、Xenium、CosMx、Slide-seq 等平台的方法学论文与技术综述，梳理了不同技术在空间分辨率（spot 级 → 单细胞级 → 亚细胞级）、检测通量、组织兼容性三个维度的差异。关键认知：当前亚细胞分辨率的主流方案为 MERFISH（基于成像）和 Stereo-seq（基于测序），这是课题"亚细胞水平共定位检测"的候选数据来源。')

add_para('（2）共定位分析方法', bold=True)
add_para('重点调研了空间转录组中共定位检测的计算方法分类：共表达网络方法（如 SpaGCN、Giotto）、空间统计方法（如 SpatialDE、SPARK-X）、基于图像的共定位像素分析方法。关键认知：现有方法大多工作在 spot 或单细胞分辨率，直接适用于亚细胞水平的共定位检测工具极度稀缺，这验证了本课题的研究缺口。')

add_para('（3）数据库与 Atlas 资源', bold=True)
add_para('梳理了 STOmicsDB、SpatialDB、SpatialOmics、TISSUE 等空间转录组数据库的收录范围、数据标准化方式与检索接口设计。关键认知：Database 类文献的标签自动化需要在"数据库名称"与"收录内容描述"之间做区分，这是平台标签策略改进的直接依据。')

add_para('（4）标签体系认知迭代', bold=True)
add_para(f'通过对 {total_articles:,} 篇文献标签分布的观察，形成了对领域热点和边缘方向的数据驱动认知。例如：Cancer ({tc.get("Cancer",0):,}) 的压倒性占比表明肿瘤是目前空间转录组学的绝对主力应用场景；"Domain Identification"标签在 2024-2025 年显著增长，反映空间域识别已成为分析流程中的标准步骤。')

doc.add_page_break()

# ── Section 3: 阶段性分析与问题识别 ──
add_heading('三、阶段性分析与问题识别', level=1)

add_heading('3.1 算法改进与效果', level=2)

add_para('本月针对标签自动化的两项核心问题完成了策略增强：')

add_para('问题1：新实体提取不足', bold=True)
add_para('Database 类和 Data Analysis 类文献中，模型容易偏向已知标签（如 "Cancer", "Visium"），导致文献标题中出现的数据库名称或方法名称未被提取。')

add_para('改进措施：', bold=True)
doc.add_paragraph('将新命名实体提取从单一冒号规则（": "前缀匹配）升级为多模式候选提取，增加正则模式覆盖 "XXX database/atlas/resource/repository" 和 "XXX method/framework/pipeline/algorithm" 等常见命名模式。', style='List Bullet')
doc.add_paragraph('引入泛词停用词表（GENERIC_NAME_STOPWORDS），过滤 "analysis", "method", "study", "pipeline" 等高频泛词，减少误报。', style='List Bullet')
doc.add_paragraph('在 Database 和 Data Analysis 类别的后处理中强化"新实体优先"策略：若成功提取新实体名称，优先使用；仅当提取失败时才回退到标签组的默认标签。', style='List Bullet')

add_para(f'目前改进已编码完成并集成到 migrate_naive.py（规则引擎）和 ml_pipeline.py（ML 链路）。标注数据规模仍然较小（{confirmed} 篇），尚不具备正式评估的条件。下月将建立专项评估集量化改进效果。')

doc.add_paragraph()

add_heading('3.2 当前模型局限性', level=2)

add_para('文本向量化：', bold=True)
add_para('当前使用 TF-IDF（词频-逆文档频率）将标题+摘要转化为稀疏向量，分类器为线性 SVM。该方案的局限在于：(1) 无法捕获上下文语义关系（如 "single-cell resolution" 和 "subcellular" 在 TF-IDF 空间中不相关，但在语义上高度关联）；(2) 对罕见词和新出现的技术名词（如新测序平台名称）缺乏泛化能力。')

add_para('主动学习的实质：', bold=True)
add_para('当前"主动学习"实际为批量被动重训：用已确认数据全量训练后预测全部未确认数据，按不确定性排序。缺少真正的 Query Strategy（如 entropy sampling 选 Top-K 最不确定的样本迭代标注），属于主动学习的简化版本。')

add_para('类别分布不均衡：', bold=True)
add_para(f'Research 类（{cat_data[0][1]:,} 篇）是 Database 类（{cat_data[4][1]} 篇）的 {cat_data[0][1]/cat_data[4][1]:.0f} 倍。极度不均衡的类别分布对 SVC 的少样本类别分类构成挑战。')

doc.add_paragraph()

add_heading('3.3 平台工程问题（已识别，待解决）', level=2)

add_para('以下工程层面的问题在本月审计中已识别并列入短期优化计划：')
doc.add_paragraph('数据并发安全：save_df() 使用全量覆盖写入（if_exists=\'replace\'），多用户快速提交时存在数据丢失风险。当前 threading.Lock 仅保护单进程。', style='List Bullet')
doc.add_paragraph('数据库表结构：literature 表未设 PRIMARY KEY，uncertainty_score 列类型为 TEXT 而非 REAL。', style='List Bullet')
doc.add_paragraph('代码重复：guess_novel_name() 等函数在 migrate_naive.py 和 ml_pipeline.py 中完全重复。', style='List Bullet')

doc.add_page_break()

# ── Section 4: 本月实验小结 ──
add_heading('四、本月实验小结', level=1)

add_para('本月主要成果：', bold=True)
doc.add_paragraph(f'搭建了可运行的 PubMed Spatial Tracker 全链路平台，入库 {total_articles:,} 篇文献，并完成 22 字段的结构化。', style='List Bullet')
doc.add_paragraph(f'完成了规则引擎与 ML 双轨推断通道的建立与约束策略对齐，{confirmed} 篇人工确认样本已可用于模型训练。', style='List Bullet')
doc.add_paragraph('通过系统性文献阅读，建立了对空间转录组技术平台、共定位分析方法、数据库资源三个维度的初步认知框架，并确认了"亚细胞水平共定位检测工具稀缺"这一研究缺口。', style='List Bullet')
doc.add_paragraph('完成了标签自动化中新实体提取的策略增强和泛词过滤机制。', style='List Bullet')

doc.add_paragraph()

add_para('主要不足与挑战：', bold=True)
doc.add_paragraph(f'人工确认比例仅 6.84%（{confirmed}/{total_articles:,}），模型训练样本严重不足，预测可靠性有限。', style='List Bullet')
doc.add_paragraph('算法改进尚未量化评估——新实体提取的召回率、首标签命中率、假阳性率均无数据支撑。', style='List Bullet')
doc.add_paragraph('目前成果以"平台搭建与知识整理"为主，尚未形成明确的生物学问题验证或方法学 benchmark。', style='List Bullet')
doc.add_paragraph('部分工程问题（并发写入、数据库表结构）虽已识别但尚未修复。', style='List Bullet')

doc.add_page_break()

# ── Section 5: 下月实验计划 ──
add_heading('五、下月实验计划', level=1)

plan_items = [
    ('建立专项评估集', 
     '针对 Database 和 Data Analysis 类别构造人工金标准（每类 ≥ 30 篇），量化新实体提取的召回率、精确率与 F1 值。同时评估类别预测的整体准确率与各类别 F1 分数，输出 ML_Performance_Report。',
     '第1-2周'),
    ('提升标注效率与确认率', 
     f'按不确定性分数排序，优先标注 Top-100 最高不确定文献；目标将人工确认率从 6.84% 提升至 10% 以上（即超过 {int(total_articles*0.1)} 篇确认样本）。重点关注 Review 和 Technology 两个高不确定性类别。',
     '持续'),
    ('修复 P0/P1 工程问题', 
     '修复 save_df() 并发写入问题（改用逐行 UPDATE + 事务），为 literature 表添加 pmid PRIMARY KEY 约束，消除代码重复。',
     '第1周'),
    ('形成小型研究问题', 
     '基于文献库数据驱动分析 + 文献阅读认知，提出一个可汇报的研究问题。候选方向示例："不同空间技术在肿瘤研究中的应用分布差异"、"亚细胞分辨率共定位检测方法的计算框架比较"。',
     '第3-4周'),
    ('输出课程汇报材料', 
     '将数据库自动生成的分析图表（年份趋势、类别分布、标签热度、期刊分布）与文献阅读结论整合为正式 PPT 汇报稿，准备科研训练中期/阶段性汇报。',
     '第4周'),
]

for i, (title, desc, timeline) in enumerate(plan_items, 1):
    add_para(f'{title}（{timeline}）', bold=True)
    add_para(desc)
    doc.add_paragraph()

# ── Signature block ──
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run('所在实验室：杨雪瑞实验室         导师签字：______________')
r.font.size = Pt(10)

# ═══════════════════ SAVE ═══════════════════
report_path = os.path.join(OUT, '月度报告-李哲夫26.04.17.docx')
doc.save(report_path)
print(f'\nReport saved to: {report_path}')
print('Done.')
